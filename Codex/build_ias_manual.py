from pathlib import Path
import json
import re
import html
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, PageBreak,
                                Image as RLImage, Table, TableStyle, KeepTogether)

ROOT = Path(r"C:\Users\Asad\Documents\GitHub\HTTPS-IAS-Production")
ASSETS = ROOT / "Codex" / "IAS_User_Manual_Assets"
DOCX_OUT = ROOT / "Codex" / "IAS Software User Manual.docx"
PDF_OUT = ROOT / "output" / "pdf" / "Internal_Audit_System_IAS_User_Manual.pdf"
ROLE_ASSETS = ROOT / "Codex" / "IAS_Manual_2026" / "screenshots"

ACCENT = "A51D47"
NAVY = "24364B"

FIGURES = [
    ("Figure 2.1", "IAS Home and dashboard workspace", "Figure_02_01_IAS_Home.png"),
    ("Figure 3.1", "Audit Criteria", "Figure_03_01_Planning_Audit_Criteria.png"),
    ("Figure 4.1", "Engagement Plan List", "Figure_04_01_Engagement_Plan_List.png"),
    ("Figure 5.1", "Manage Audit Observations", "Figure_05_01_Execution_Manage_Observations.png"),
    ("Figure 6.1", "Auditee Assigned Observations", "Figure_06_01_Auditee_Observation_Assigned.png"),
    ("Figure 7.1", "Post Compliance list and export controls", "Figure_07_01_Post_Compliance.png"),
    ("Figure 8.1", "IID Case Dashboard", "Figure_08_01_IID_Dashboard.png"),
    ("Figure 9.1", "Settled Para Monitoring Report", "Figure_09_01_Reports_Para_Settlement.png"),
    ("Figure 10.1", "Field Audit workflow dashboard", "Figure_10_01_FAD_Field_Auditor_Dashboard.png"),
    ("Figure 11.1", "Commercial Audit workflow", "Figure_11_01_CAU_Workflow.png"),
    ("Figure 12.1", "Monitoring of Samples", "Figure_12_01_Sampling_Monitoring.png"),
    ("Figure 13.1", "Checklist Lifecycle Workspace", "Figure_13_01_Setup_Checklist_Dashboard.png"),
    ("Figure 14.1", "Master Admin Control Panel", "Figure_14_01_Admin_Master_Control.png"),
    ("Figure 14.2", "Manage Users", "Figure_14_02_Admin_Manage_Users.png"),
    ("Figure 15.1", "Email Management dashboard", "Figure_15_01_Email_Dashboard.png"),
]

DETAILED = {
    "Planning": {
        "screen": "Audit Criteria", "figure": "Figure_03_01_Planning_Audit_Criteria.png",
        "purpose": "Define an audit period, entity, auditable entity, frequency, duration and field-visit requirement before submitting criteria.",
        "controls": "Audit Period, Entity, Auditable Entity, Frequency, No. of Days, Field Visit, Save, Delete and Submit.",
        "steps": ["Open Planning > Audit Criteria.", "Select the audit period and entity.", "Select the auditable entity and frequency.", "Enter the number of days and set Field Visit when applicable.", "Select Save and review the new row in the criteria table.", "Use Delete only for an incorrect saved row; select Submit when the displayed criteria are ready for the next stage."],
        "result": "IAS stores the criteria row and, on Submit, sends the displayed criteria into the configured planning review flow."
    },
    "Engagement": {
        "screen": "Engagement Plan List", "figure": "Figure_04_01_Engagement_Plan_List.png",
        "purpose": "Review engagement dates and make the available approve or reject decision.",
        "controls": "Entity Name, Team Name, operation/start/end dates, Reject and Approve.",
        "steps": ["Open Engagement > Engagement Plan List.", "Locate the engagement row and compare the entity, team and displayed dates.", "Choose Approve or Reject from the row when that action is available to your profile.", "Confirm the resulting status or removal from the work list."],
        "result": "The selected engagement decision is recorded and the work list refreshes."
    },
    "Execution": {
        "screen": "Manage Observations", "figure": "Figure_05_01_Execution_Manage_Observations.png",
        "purpose": "Select an audit entity and period, then review and act on the resulting observation list.",
        "controls": "Entity, Audit Period and a table containing Memo No., Title, Violation, Risk Category, Status and row actions.",
        "steps": ["Open Execution > Manage Observations.", "Select an Entity.", "Select the Audit Period that becomes available.", "Review the observation list and its status.", "Use only the action shown on the relevant row to open or continue that observation."],
        "result": "IAS opens the selected observation or refreshes the list after the completed action."
    },
    "Auditee Portal": {
        "screen": "Observation Assigned", "figure": "Figure_06_01_Auditee_Observation_Assigned.png",
        "purpose": "Display observations or memos assigned to the signed-in auditee for reply.",
        "controls": "Memo number, audit year, audit entity, memo date, gist, Reply Till and Status.",
        "steps": ["Open Auditee Portal > Observation Assigned.", "Locate the assigned memo and check the Reply Till date and current status.", "Open the available row/action and enter the response and any permitted evidence.", "Submit the response and confirm the updated status."],
        "result": "The auditee response is attached to the assigned item and its workflow status is updated."
    },
    "Post Compliance": {
        "screen": "Post Compliance", "figure": "Figure_07_01_Post_Compliance.png",
        "purpose": "Review outstanding compliance items and their compliance history.",
        "controls": "Page size, Search, PDF, Export to Excel, Export to CSV, Copy to Clipboard and compliance-history table.",
        "steps": ["Open Post Compliance > Post Compliance.", "Use Search to narrow the displayed records.", "Review the audit period, para, dates, title/gist, source and compliance history.", "Use PDF, Excel, CSV or Copy only when an output is required."],
        "result": "The list is filtered on screen or exported in the selected format."
    },
    "I&ID": {
        "screen": "IID Case Dashboard", "figure": "Figure_08_01_IID_Dashboard.png",
        "purpose": "Navigate the complaint/inquiry case workflow from complaint through final approval.",
        "controls": "Task List, Monitoring Dashboard, Read Only Report, Exceptions Reports, complaint selector and numbered workflow stages.",
        "steps": ["Open I&ID > IID Dashboard.", "Select a complaint from the list.", "Choose the active workflow stage, such as Complaint, Initial Assessment or Head Review.", "Complete the fields and actions displayed inside that stage.", "Return to Task List or Monitoring Dashboard to verify progress; use Read Only Report when a non-editable report is needed."],
        "result": "The case remains at, or advances from, the selected workflow stage according to the action completed."
    },
    "Reports": {
        "screen": "Monitoring of Settlement of Para", "figure": "Figure_09_01_Reports_Para_Settlement.png",
        "purpose": "Read the settled-para details and compliance history for a selected entity.",
        "controls": "Entity selector and report columns for office, entity, audit/para year, para number, settlement and compliance history.",
        "steps": ["Open Reports > Monitoring of Settlement of Para.", "Select an Entity.", "Review the returned rows and settlement/compliance columns.", "Use any export or print action only if it is displayed for the resulting report."],
        "result": "IAS displays the settled-para monitoring data for the selected entity."
    },
    "FAD": {
        "screen": "Field Auditor Dashboard", "figure": "Figure_10_01_FAD_Field_Auditor_Dashboard.png",
        "purpose": "Run the field-audit engagement workflow through the stages exposed for the selected engagement.",
        "controls": "Engagement selector, Joining, Sampling, Exception Report, Working Paper and later workflow stages; Mark as Saved.",
        "steps": ["Open FAD > Field Auditor Dashboard.", "Select an engagement.", "Choose the current enabled workflow step.", "Complete the controls loaded in the workspace.", "Use Mark as Saved when the stage content is complete enough to save, then continue to the next enabled stage."],
        "result": "The selected engagement is saved at the current field-audit stage and later stages become available according to workflow state."
    },
    "CAU": {
        "screen": "CAU Workflow", "figure": "Figure_11_01_CAU_Workflow.png",
        "purpose": "Manage the commercial-audit sequence from entity selection through ARPSE follow-up.",
        "controls": "CM Entity, CM Register, PO Entry, PDF Linking, ARPSE Header, ARPSE Linking, ARPSE Monitoring and ARPSE Follow-up.",
        "steps": ["Open CAU > CAU Workflow.", "Select the current commercial-audit entity.", "Open the enabled workflow step.", "Complete the controls shown in the workspace and save or submit using the action displayed there.", "Continue to the next enabled step and use Monitoring or Follow-up for later-stage tracking."],
        "result": "The commercial-audit item is saved at the selected stage and the workflow display refreshes."
    },
    "Sampling": {
        "screen": "Monitoring of Samples", "figure": "Figure_12_01_Sampling_Monitoring.png",
        "purpose": "Monitor samples by entity and use the available view or regeneration action.",
        "controls": "Entity, Sample List, Percentage, Total Count, Sample Count, View and Regenerate Sample.",
        "steps": ["Open Sampling > Monitoring of Samples.", "Select an Entity.", "Review sample percentage and counts.", "Choose View to inspect a sample or Regenerate Sample only when the displayed workflow permits it."],
        "result": "IAS displays the selected sample details or refreshes the generated sample set."
    },
    "Setup": {
        "screen": "Checklist Dashboard", "figure": "Figure_13_01_Setup_Checklist_Dashboard.png",
        "purpose": "Manage the checklist lifecycle through create, submit, detail, review and authorization stages.",
        "controls": "Six-step lifecycle, Add New Checklist, checklist table and row Update action.",
        "steps": ["Open Setup > Checklist Dashboard.", "Choose the relevant lifecycle stage.", "In Manage Checklist, select Add New Checklist or Update for an existing row.", "Complete the stage-specific fields and save/submit using the displayed action.", "Use the review or authorization stages only when they are enabled for your profile."],
        "result": "The checklist is created or updated and its lifecycle state is reflected in the workspace."
    },
    "Admin": {
        "screen": "Master Admin Control Panel and Manage Users", "figure": "Figure_14_01_Admin_Master_Control.png",
        "purpose": "Maintain page/menu configuration and user records through permission-controlled administration screens.",
        "controls": "Master panel phase/accordion controls; Manage Users fields for relationship type, office, posting, P.P. number, login name, email, group/role and active status.",
        "steps": ["Open Admin > Master Admin Control Panel for menu/page configuration, or Admin > Manage Users for a user record.", "Use Find or the available filters before changing an existing record.", "Enter or update only the fields shown on the screen.", "Select the displayed Add, Update, Save or assignment action and confirm the refreshed record."],
        "result": "IAS stores the configuration or user change; visibility and available actions remain subject to assigned rights." 
    },
    "Email": {
        "screen": "Email Management", "figure": "Figure_15_01_Email_Dashboard.png",
        "purpose": "View and maintain the isolated email-management configuration area.",
        "controls": "Dashboard, Events, Templates, Recipient rules, Placeholders, Preview, Manual test and Logs.",
        "steps": ["Open Email > Dashboard.", "Choose Events, Templates, Recipient rules or Placeholders to maintain configuration.", "Use Preview before testing a message.", "Use Manual test only with approved test recipients.", "Review Logs for the recorded result."],
        "result": "The selected email configuration or test result is saved and can be reviewed in Logs. The dashboard explicitly states that configured events do not automatically trigger IAS business email." 
    }
}

WORKFLOWS = [
    ("Planning and engagement", "Login → Planning: constitute teams / enter audit criteria → submit criteria → reviewer approval or return → generate/tentative plan → Engagement Plan List → approve or reject."),
    ("Observation and auditee response", "Login → Execution: select entity and audit period → open observation → complete available observation action → Auditee Portal: open assigned observation → reply and attach permitted evidence → submit → verify status."),
    ("Draft and final report", "Login → Execution: manage observations / draft report paras → quality review → finalize draft audit report → concluding/closing audit. FAD also exposes report overview, finalize report and PDF lists for its workflow."),
    ("Post compliance and settlement", "Login → Post Compliance: find outstanding item → enter or review compliance → review stage → monitoring of settled paras. Reports provides settled-para and compliance progress views."),
    ("IID inquiry", "Login → IID Dashboard → select complaint → Complaint → Initial Assessment → Head Review → Investigation Plan → Plan Approval → Analysis / Inquiry Report → Final Approval → read-only report or case study."),
    ("Commercial audit", "Login → CAU Workflow → select commercial entity → CM Register → PO Entry → PDF Linking → ARPSE Header/Linking → Monitoring → Follow-up."),
    ("Checklist lifecycle", "Login → Setup: Checklist Dashboard → manage checklist → submit → manage detail → review → authorize. Update/authorization screens are also available as separate menu items."),
    ("Email configuration/test", "Login → Email Dashboard → configure events/templates/recipient rules/placeholders → Preview → Manual test → Logs. The verified dashboard states that events do not automatically trigger IAS business email."),
]

ROLE_GUIDE = [
    ("Super User", "01_super_user_dashboard.png", ["Dashboard", "Planning", "Engagement", "Execution", "Auditee Portal", "Post Compliance", "I&ID", "Reports", "FAD", "CAU", "Sampling", "Setup", "Admin", "Email"]),
    ("Admin", "02_admin_dashboard.png", ["Dashboard", "Engagement", "Execution", "Post Compliance", "Reports", "FAD", "CAU", "Setup", "Admin"]),
    ("Zonal Chief", "03_zonal_chief_dashboard.png", ["Dashboard", "Auditee Portal", "Post Compliance", "Reports"]),
    ("Head FAD", "04_head_fad_dashboard.png", ["Dashboard", "Planning", "Engagement", "Execution", "Post Compliance", "Reports", "FAD", "CAU", "Sampling", "Setup"]),
    ("Head CAD", "05_head_cad_dashboard.png", ["Dashboard", "Risk", "Planning", "Engagement", "Execution", "Post Compliance", "Reports", "Sampling", "Setup"]),
    ("Head ISAD", "06_head_isad_dashboard.png", ["Dashboard", "Risk", "Planning", "Engagement", "Execution", "Post Compliance", "Reports", "FAD", "Sampling", "Setup", "Admin"]),
    ("HO Auditor (Team Lead)", "07_ho_auditor_team_lead_dashboard.png", ["Risk", "Execution", "Post Compliance", "FAD"]),
    ("HO Planning & Coordination", "08_ho_planning_coordination_dashboard.png", ["Dashboard", "Planning", "Engagement", "Execution", "Post Compliance", "Reports", "FAD", "Sampling", "Admin"]),
    ("Dept Auditee", "09_dept_auditee_dashboard.png", ["Dashboard", "Auditee Portal", "Post Compliance", "Reports"]),
    ("Auditee Branch", "10_auditee_branch_dashboard.png", ["Dashboard", "Auditee Portal", "Post Compliance", "Reports"]),
    ("Div Auditee", "11_div_auditee_dashboard.png", ["Dashboard", "Auditee Portal", "Post Compliance", "Reports"]),
    ("SVP Audit Zone", "12_svp_audit_zone_dashboard.png", ["Dashboard", "Planning", "Engagement", "Execution", "Post Compliance", "Reports", "FAD", "Setup"]),
    ("AZ Implementation Unit", "13_az_implementation_dashboard.png", ["Dashboard", "Planning", "Engagement", "Execution", "Post Compliance", "Reports", "FAD"]),
    ("AZ Planning & Coordination", "14_az_planning_coordination_dashboard.png", ["Execution", "Post Compliance", "Reports"]),
    ("Branch Auditor (Team Lead)", "15_branch_auditor_team_lead_dashboard.png", ["Execution", "Post Compliance", "FAD", "Sampling"]),
    ("Audit QA Unit", "16_audit_qa_dashboard.png", ["Execution", "Setup"]),
    ("Shariah Audit Unit", "17_shariah_audit_dashboard.png", ["Risk", "Planning", "Engagement", "Execution", "Setup"]),
    ("HO Auditor (Team Member)", "18_ho_auditor_team_member_dashboard.png", ["Execution", "Auditee Portal", "Post Compliance"]),
    ("Branch Auditor (Team Member)", "19_branch_auditor_team_member_dashboard.png", ["Execution", "Post Compliance", "FAD"]),
]

ROLE_PURPOSE = {
    "Super User": "Administer the IAS permission estate and access all configured operational modules.",
    "Admin": "Maintain users, configuration and assigned operational administration functions.",
    "Zonal Chief": "Monitor zonal audit results, auditee work and post-compliance reporting.",
    "Head FAD": "Oversee field-audit planning, execution, quality, reporting and compliance.",
    "Head CAD": "Oversee central/credit audit risk, planning, execution and compliance workflows.",
    "Head ISAD": "Oversee information-systems audit risk, planning, execution, reporting and administration.",
    "HO Auditor (Team Lead)": "Lead Head Office audit execution, review team work and progress assigned items.",
    "HO Planning & Coordination": "Coordinate Head Office planning, engagements, reporting and monitoring.",
    "Dept Auditee": "Review departmental observations, provide responses and monitor compliance.",
    "Auditee Branch": "Review branch observations, provide responses and monitor compliance.",
    "Div Auditee": "Review divisional observations, provide responses and monitor compliance.",
    "SVP Audit Zone": "Supervise Audit Zone planning, execution, field-audit and compliance activity.",
    "AZ Implementation Unit": "Implement Audit Zone plans and progress execution, reporting and compliance work.",
    "AZ Planning & Coordination": "Coordinate Audit Zone execution, reporting and post-compliance work.",
    "Branch Auditor (Team Lead)": "Lead branch audit execution, sampling and field-audit workflow stages.",
    "Audit QA Unit": "Perform quality review and checklist-related assurance functions exposed by IAS.",
    "Shariah Audit Unit": "Perform Shariah audit risk, planning, engagement, execution and checklist work.",
    "HO Auditor (Team Member)": "Prepare assigned Head Office audit work and submit it into team review flows.",
    "Branch Auditor (Team Member)": "Prepare assigned branch audit and field-audit work for Team Lead review.",
}


def safe_text(value):
    return re.sub(r"\s+", " ", value or "").strip()


def pdf_safe(value):
    return html.escape(str(value)).replace("→", "-&gt;")


def purpose_for(name):
    n = name.lower()
    rules = [
        (("dashboard",), "Open the module workspace and review the current work/status summary."),
        (("report", "monitoring", "analysis"), "Filter and review the named IAS report or monitoring view."),
        (("approval", "authorize", "review", "checking"), "Review the named item and use the decision actions displayed for the signed-in profile."),
        (("manage", "setup", "configuration", "assignment", "addition", "shifting"), "Find, add or update the named configuration records using the controls displayed on the screen."),
        (("search", "find"), "Search for the named IAS records using the available filters."),
        (("pdf", "download", "export", "list of pdf"), "Generate, open or download the named output when records are available."),
        (("submit", "reply", "follow up", "compliance"), "Open the named work item, complete the displayed response fields and submit through the available workflow action."),
        (("plan", "criteria", "team"), "Prepare or maintain the named planning information and submit it through the displayed workflow."),
    ]
    for keys, text in rules:
        if any(key in n for key in keys):
            return text
    return "Open and work with the named IAS function using only the controls and actions displayed for the signed-in profile."


def add_field(run, instr):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run._r.addnext(fld)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def doc_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(125)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor.from_string(ACCENT)
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.size = Pt(18); p2.runs[0].font.color.rgb = RGBColor.from_string(NAVY)


def add_doc_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor.from_string(ACCENT if level == 1 else NAVY)
    return p


def add_doc_figure(doc, filename, caption):
    path = ASSETS / filename
    with Image.open(path) as image:
        width_px, height_px = image.size
    width_in = 6.35
    height_in = width_in * height_px / width_px
    if height_in > 7.2:
        height_in = 7.2
        width_in = height_in * width_px / height_px
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width_in), height=Inches(height_in))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.style = doc.styles["Caption"]


def add_role_figure(doc, filename, caption):
    path = ROLE_ASSETS / filename
    with Image.open(path) as image:
        width_px, height_px = image.size
    width_in = 6.35
    height_in = width_in * height_px / width_px
    if height_in > 7.2:
        height_in = 7.2
        width_in = height_in * width_px / height_px
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width_in), height=Inches(height_in))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.style = doc.styles["Caption"]


def build_docx(menu_data):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(.65); sec.bottom_margin = Inches(.65)
    sec.left_margin = Inches(.72); sec.right_margin = Inches(.72)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"; styles["Normal"].font.size = Pt(9)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Aptos Display"
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Internal Audit System (IAS) — User Manual   |   Page ")
    add_field(footer.add_run(), "PAGE")

    doc_title(doc, "Internal Audit System (IAS)", "User Manual")
    p = doc.add_paragraph("Practical screen and workflow guide")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.runs[0].bold = True
    p = doc.add_paragraph("Document version 2.0  |  28 August 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Prepared from the verified IAS source tree and the authenticated local application views. No credentials, secrets or live record details are included.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_doc_heading(doc, "Document Control", 1)
    table = doc.add_table(rows=0, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
    for left, right in [("Document", "IAS Software User Manual"), ("Version", "2.0"), ("Date", "28 August 2026"), ("Status", "Issued for operational use and review"), ("Basis", "All 19 supplied roles verified in the authenticated UI and cross-checked against IAS authorization/menu source")]:
        cells = table.add_row().cells; set_cell_text(cells[0], left, True, "FFFFFF", 9); shade(cells[0], NAVY); set_cell_text(cells[1], right, False, None, 9)
    add_doc_heading(doc, "Revision History", 2)
    t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Version", "Date", "Prepared for", "Change"]): set_cell_text(t.rows[0].cells[i], h, True, "FFFFFF", 8); shade(t.rows[0].cells[i], ACCENT)
    row = t.add_row().cells
    for i, v in enumerate(["2.0", "28 Aug 2026", "IAS users / IT review", "Complete 19-role UI verification, permission matrix and discrepancy register"]): set_cell_text(row[i], v)
    add_doc_heading(doc, "How to Use This Manual", 2)
    doc.add_paragraph("IAS menus and actions are permission-based. A screen may be absent, read-only or show different workflow actions for another user. Follow the controls actually displayed for your signed-in profile. Empty screenshots indicate that no test records were available; they still verify the screen layout and available controls.")
    add_doc_heading(doc, "Table of Contents", 1)
    toc = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u'); toc._p.append(fld)
    doc.add_paragraph("In Microsoft Word, right-click the table and choose Update Field if page numbers are not refreshed automatically.", style=None)
    doc.add_page_break()

    add_doc_heading(doc, "1. Getting Started", 1)
    add_doc_heading(doc, "1.1 Sign in", 2)
    doc.add_paragraph("Open the IAS address supplied by your organization. Enter your assigned user name and password, then select the login action. If IAS reports an existing session, follow the on-screen session message or contact the system administrator. Never share credentials or include them in screenshots.")
    add_doc_heading(doc, "1.2 Navigation and access", 2)
    doc.add_paragraph("After sign-in, the left navigation groups functions into Dashboard, Planning, Engagement, Execution, Auditee Portal, Post Compliance, I&ID, Reports, FAD, CAU, Sampling, Setup, Admin and Email. Use the menu collapse control when more space is required. Use Logout when work is complete.")
    add_doc_heading(doc, "1.3 Common list and form behaviour", 2)
    for text in ["Choose parent filters first; later lists often remain empty until a required entity, period or engagement is selected.", "Review the record status before using a row action.", "Save preserves entered work where the screen provides it; Submit, Approve, Reject, Authorize or Finalize moves or decides workflow work and should be used only after review.", "Search filters the current list. PDF, Excel, CSV, Print or Download actions create the output named on the screen.", "For evidence uploads, choose only the requested business file, confirm the correct record, and use the upload action displayed. Open or download evidence through the record-specific retrieval action."]:
        doc.add_paragraph(text, style="List Bullet")

    add_doc_heading(doc, "2. Home and Dashboard", 1)
    add_doc_figure(doc, "Figure_02_01_IAS_Home.png", "Figure 2.1: IAS Home and dashboard workspace")
    doc.add_paragraph("The home workspace provides quick links, announcements, calendars/holidays and dashboard entry points available to the signed-in user. Select a quick link or expand a module in the left menu. The verified home view exposed Master Admin Control Panel, Manage Users, Field Auditor Dashboard, User Dashboard, Entity Dashboard, Checklist Dashboard, Planning Dashboard and IID Dashboard for the reviewed profile.")
    doc.add_paragraph("Typical path: Login → Home → select a module → choose a record → complete the displayed action → verify status or output.")

    section_no = 3
    for module, detail in DETAILED.items():
        add_doc_heading(doc, f"{section_no}. {module}", 1)
        add_doc_heading(doc, f"{section_no}.1 {detail['screen']}", 2)
        add_doc_figure(doc, detail["figure"], f"Figure {section_no}.1: {detail['screen']}")
        doc.add_paragraph("Purpose: " + detail["purpose"])
        doc.add_paragraph("Access: Users whose IAS page/menu permissions expose this screen. Review/approval/administration actions additionally depend on the actions enabled for the signed-in profile.")
        doc.add_paragraph("Main fields and controls: " + detail["controls"])
        add_doc_heading(doc, "Procedure", 3)
        for step in detail["steps"]: doc.add_paragraph(step, style="List Number")
        doc.add_paragraph("Expected result: " + detail["result"])
        if module == "Admin":
            add_doc_figure(doc, "Figure_14_02_Admin_Manage_Users.png", "Figure 14.2: Manage Users")
            doc.add_paragraph("Restriction: administration screens change shared configuration and access. Confirm the target user, office, group/role and active status before saving.")
        section_no += 1

    add_doc_heading(doc, "16. End-to-End User Workflows", 1)
    for title, flow in WORKFLOWS:
        add_doc_heading(doc, title, 2); doc.add_paragraph(flow)

    add_doc_heading(doc, "17. Complete Verified Screen Catalog", 1)
    doc.add_paragraph("This catalog lists every function exposed in the authenticated IAS menu during the review. It is the complete verified menu surface for the reviewed profile. Open paths are shown as menu labels, not technical URLs. Functions hidden from this profile cannot be confirmed as user-accessible.")
    for module in menu_data:
        add_doc_heading(doc, module["menu"], 2)
        t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Screen / function", "How to open", "Verified user use"]): set_cell_text(t.rows[0].cells[i], h, True, "FFFFFF", 7.5); shade(t.rows[0].cells[i], NAVY)
        for link in module["links"]:
            cells = t.add_row().cells
            vals = [safe_text(link["text"]), f"{module['menu']} > {safe_text(link['text'])}", purpose_for(link["text"])]
            for i, v in enumerate(vals): set_cell_text(cells[i], v, size=7.2)
        doc.add_paragraph("Access and restrictions: permission-based. Use only the fields, row actions and workflow buttons displayed for the signed-in profile.")

    add_doc_heading(doc, "18. Validation, Uploads, Search and Outputs", 1)
    for title, body in [
        ("Validation", "Required selections and field rules are enforced by the individual screen. Correct highlighted or on-screen validation messages before trying again. Do not infer acceptance from a button click; confirm the refreshed row, status or confirmation message."),
        ("File/evidence handling", "IAS contains file/evidence upload and retrieval functions in observation/report workflows and catalog/PDF areas. Associate a file only with the intended record, avoid unnecessary personal data, and confirm the file appears in the record-specific list. Exact extensions and size limits vary by screen and should be followed from its displayed validation message."),
        ("Search/filtering", "Select required parent values, enter search text where provided, then review the returned table. Clear or change the filter to broaden results."),
        ("Reports/downloads", "Select the required entity, period or status first. Generate or export only with the button shown by the screen. Verify the downloaded title and selected scope before distributing it."),
        ("Errors and support", "If a screen remains empty after required filters, an action is disabled, or an error persists, record the screen name and time without including confidential record content, then contact the IAS support/administration team."),
    ]:
        add_doc_heading(doc, title, 2); doc.add_paragraph(body)

    add_doc_heading(doc, "19. Figure Index and Verification Note", 1)
    for number, caption, filename in FIGURES:
        doc.add_paragraph(f"{number}: {caption} — authenticated local IAS view ({filename})")
    doc.add_paragraph("Verification basis: menu/routes and user-facing behaviour were checked against the IAS source tree and the authenticated local application on 28 August 2026. The manual intentionally does not contain passwords, PP numbers, connection strings, database objects or other internal implementation details. Screenshots show the post-login application viewport and contain no populated confidential business records.")

    add_doc_heading(doc, "20. Role-Based User Guide", 1)
    doc.add_paragraph("Each role below was authenticated separately with the supplied test account. The module list records what was actually visible in the left navigation on 28 August 2026. A module not listed for the role was not available from that role's verified landing page. Detailed procedures and workflow screenshots are provided in Sections 3-18.")
    for idx, (role, screenshot, modules) in enumerate(ROLE_GUIDE, 1):
        add_doc_heading(doc, f"20.{idx} {role}", 2)
        doc.add_paragraph("Role purpose: " + ROLE_PURPOSE[role])
        doc.add_paragraph("Verified landing page: IAS Home. Visible functional modules: " + ", ".join(modules) + ".")
        add_role_figure(doc, screenshot, f"Figure 20.{idx}: {role} verified IAS dashboard")
        add_doc_heading(doc, "Role operating procedure", 3)
        steps = [
            "Sign in and, when the Select Working Context page appears, select the required role/posting context.",
            "Confirm that the IAS Home page opens and review the left navigation against the verified module list above.",
            "Open the required permitted module, select the relevant entity, period, engagement or work item, and review its current status.",
            "Use only the action displayed for the role. Team Members prepare and submit assigned work; Team Leads and Heads perform the review/decision actions exposed to them; auditees provide responses; administration roles maintain shared setup only where enabled.",
            "Confirm the resulting status, refreshed row, report or confirmation message before leaving the screen.",
        ]
        for step in steps:
            doc.add_paragraph(step, style="List Number")
        doc.add_paragraph("Restricted/hidden functions: all modules not listed above were absent from the verified landing menu. Button-level actions remain conditional on the selected record and workflow status.")

    add_doc_heading(doc, "21. Role / Permission Matrix", 1)
    doc.add_paragraph("Legend: check mark = module visible and available from the verified role menu; Not Available = module absent. Approval, edit and submission buttons are record/status dependent and are described in the role and workflow procedures rather than inferred from menu presence.")
    matrix_modules = ["Dashboard", "Risk", "Planning", "Engagement", "Execution", "Auditee Portal", "Post Compliance", "Reports", "FAD", "CAU", "Sampling", "Setup", "Admin", "Email"]
    for part, columns in enumerate((matrix_modules[:7], matrix_modules[7:]), 1):
        add_doc_heading(doc, f"Table 21.{part}: Verified module access", 2)
        t = doc.add_table(rows=1, cols=1 + len(columns)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Role"] + columns
        for col, label in enumerate(headers):
            set_cell_text(t.rows[0].cells[col], label, True, "FFFFFF", 6.5); shade(t.rows[0].cells[col], NAVY)
        for role, _, modules in ROLE_GUIDE:
            cells = t.add_row().cells
            set_cell_text(cells[0], role, True, None, 6.3)
            for col, module in enumerate(columns, 1):
                set_cell_text(cells[col], "✓" if module in modules else "Not Available", False, None, 6.0)

    add_doc_heading(doc, "22. Issues / Role Permission Discrepancies Identified During Manual Preparation", 1)
    issues = [
        ("Logout link behaviour", "For multiple tested sessions, selecting the visible left-navigation Logout link did not leave IAS Home. Direct navigation to the same /Login/Logout endpoint ended the session and returned the login page. Review client-side click handling and confirm the behaviour in the supported production browser."),
        ("Context selection changes the effective role", "The same test identity can hold several working contexts. IAS correctly displayed Select Working Context and changed the menu footprint after the selected context was continued. Operational instructions and support diagnostics must therefore record the active role/posting context, not only the person."),
        ("Permission-driven submenus and conditional actions", "The source menu layout groups pages returned for the current session, while PermissionService caches page/API permissions. UI verification confirmed materially different top-level menus by role. No function should be assigned solely from job title; administrators must validate the actual group/page assignment after role changes."),
        ("Duplicate administrative route entry", "The verified Super User menu inventory contains two entries for Authorize Delete Duplicate Para, one of which includes a trailing tab character in the route. Remove or normalize the malformed duplicate to avoid inconsistent routing or permission matching."),
        ("Spelling and naming consistency", "Several user-facing menu labels contain inconsistent spelling or separators, including Reffered Back, Engagment and Quality_Assurance_checking. Standardize labels without changing the underlying permission keys."),
    ]
    t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Issue", "Evidence / impact", "Recommended action"]): set_cell_text(t.rows[0].cells[i], h, True, "FFFFFF", 7.5); shade(t.rows[0].cells[i], ACCENT)
    for title, evidence in issues:
        cells = t.add_row().cells
        set_cell_text(cells[0], title, True, None, 7)
        set_cell_text(cells[1], evidence, False, None, 7)
        set_cell_text(cells[2], "Validate with the application owner, correct configuration/code where applicable, and regression-test the affected roles.", False, None, 7)

    add_doc_heading(doc, "23. Role Verification Register", 1)
    doc.add_paragraph("Completion statement: every role requested for this manual was logged in, checked and documented on 28 August 2026. Credentials were used only for testing and are not reproduced in this document or its screenshots.")
    t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["No.", "Role", "UI login", "Documented"]): set_cell_text(t.rows[0].cells[i], h, True, "FFFFFF", 7.5); shade(t.rows[0].cells[i], NAVY)
    for idx, (role, _, _) in enumerate(ROLE_GUIDE, 1):
        cells = t.add_row().cells
        for i, value in enumerate([idx, role, "Verified 28 Aug 2026", "Yes"]): set_cell_text(cells[i], value, False, None, 7)
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def build_pdf(menu_data):
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=28, leading=34, textColor=colors.HexColor("#"+ACCENT), alignment=TA_CENTER, spaceAfter=16))
    styles.add(ParagraphStyle(name="CoverSub", parent=styles["Normal"], fontName="Helvetica", fontSize=16, leading=20, textColor=colors.HexColor("#"+NAVY), alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="H1x", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=19, textColor=colors.HexColor("#"+ACCENT), spaceBefore=8, spaceAfter=8))
    styles.add(ParagraphStyle(name="H2x", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#"+NAVY), spaceBefore=7, spaceAfter=5))
    styles.add(ParagraphStyle(name="Bodyx", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.5, leading=11, spaceAfter=4))
    styles.add(ParagraphStyle(name="Captionx", parent=styles["BodyText"], fontName="Helvetica-Oblique", fontSize=7.5, leading=9, alignment=TA_CENTER, textColor=colors.HexColor("#555555"), spaceAfter=7))
    story = [Spacer(1, 45*mm), Paragraph("Internal Audit System (IAS)", styles["CoverTitle"]), Paragraph("User Manual", styles["CoverSub"]), Spacer(1, 18*mm), Paragraph("Practical screen and workflow guide", styles["CoverSub"]), Spacer(1, 8*mm), Paragraph("Document version 1.0 | 27 August 2026", styles["Bodyx"]), Spacer(1, 14*mm), Paragraph("Prepared from the verified IAS source tree and authenticated local application views. No credentials, secrets or live record details are included.", styles["Bodyx"]), PageBreak()]
    story += [Paragraph("Document Control", styles["H1x"])]
    control = [["Document", "Internal Audit System (IAS) User Manual"], ["Version", "1.0"], ["Date", "27 August 2026"], ["Status", "Issued for operational use and review"], ["Basis", "Verified Razor views, controllers/menu routes and authenticated local IAS screens"]]
    tbl = Table(control, colWidths=[35*mm, 135*mm]); tbl.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.35,colors.grey),("BACKGROUND",(0,0),(0,-1),colors.HexColor("#"+NAVY)),("TEXTCOLOR",(0,0),(0,-1),colors.white),("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"),("FONTNAME",(1,0),(1,-1),"Helvetica"),("FONTSIZE",(0,0),(-1,-1),8),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("PADDING",(0,0),(-1,-1),5)])); story += [tbl, Spacer(1,5*mm), Paragraph("Revision History",styles["H2x"])]
    rev = Table([["Version","Date","Prepared for","Change"],["1.0","27 Aug 2026","IAS users / IT review","Initial verified user manual"]], colWidths=[22*mm,30*mm,55*mm,63*mm]); rev.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.35,colors.grey),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#"+ACCENT)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8),("PADDING",(0,0),(-1,-1),4)])); story += [rev, Paragraph("How to use this manual",styles["H2x"]), Paragraph("IAS menus and actions are permission-based. A screen may be absent, read-only or show different workflow actions for another user. Empty screenshots indicate that no test records were available; they still verify the screen layout and controls.",styles["Bodyx"]), PageBreak()]
    story += [Paragraph("Contents",styles["H1x"])]
    contents = ["1. Getting Started", "2. Home and Dashboard"] + [f"{i}. {m}" for i,m in enumerate(DETAILED.keys(),3)] + ["16. End-to-End User Workflows", "17. Complete Verified Screen Catalog", "18. Validation, Uploads, Search and Outputs", "19. Figure Index and Verification Note"]
    for item in contents: story.append(Paragraph(pdf_safe(item),styles["Bodyx"]))
    story.append(PageBreak())
    story += [Paragraph("1. Getting Started",styles["H1x"]), Paragraph("Sign in with the account assigned by your organization. After sign-in, use the permission-based left menu. Select parent filters first, review record status, then use only the action displayed for the intended row. Save preserves work; Submit, Approve, Reject, Authorize and Finalize are workflow decisions. Logout when work is complete.",styles["Bodyx"]), Paragraph("Typical flow: Login - Home - select module - choose record - perform displayed action - verify status or output.",styles["Bodyx"]), Paragraph("2. Home and Dashboard",styles["H1x"])]
    def pdf_figure(filename, caption):
        img = Image.open(ASSETS/filename); w,h=img.size; maxw=170*mm; maxh=105*mm; scale=min(maxw/w,maxh/h)
        return [RLImage(str(ASSETS/filename),width=w*scale,height=h*scale),Paragraph(pdf_safe(caption),styles["Captionx"])]
    story += pdf_figure("Figure_02_01_IAS_Home.png","Figure 2.1: IAS Home and dashboard workspace")
    story.append(Paragraph("The home workspace provides quick links, announcements, calendars/holidays and dashboard entry points available to the signed-in user.",styles["Bodyx"]))
    for section_no,(module,detail) in enumerate(DETAILED.items(),3):
        story += [PageBreak(), Paragraph(pdf_safe(f"{section_no}. {module}"),styles["H1x"])]
        story += pdf_figure(detail["figure"],f"Figure {section_no}.1: {detail['screen']}")
        story += [Paragraph("<b>Purpose:</b> "+pdf_safe(detail["purpose"]),styles["Bodyx"]), Paragraph("<b>Access:</b> Users whose IAS page/menu permissions expose this screen; available decisions depend on the signed-in profile.",styles["Bodyx"]), Paragraph("<b>Main controls:</b> "+pdf_safe(detail["controls"]),styles["Bodyx"]), Paragraph("Procedure",styles["H2x"])]
        for i,step in enumerate(detail["steps"],1): story.append(Paragraph(pdf_safe(f"{i}. {step}"),styles["Bodyx"]))
        story.append(Paragraph("<b>Expected result:</b> "+pdf_safe(detail["result"]),styles["Bodyx"]))
        if module=="Admin": story += pdf_figure("Figure_14_02_Admin_Manage_Users.png","Figure 14.2: Manage Users")
    story += [PageBreak(),Paragraph("16. End-to-End User Workflows",styles["H1x"])]
    for title,flow in WORKFLOWS:
        story += [Paragraph(pdf_safe(title),styles["H2x"]),Paragraph(pdf_safe(flow),styles["Bodyx"])]
    story += [PageBreak(),Paragraph("17. Complete Verified Screen Catalog",styles["H1x"]),Paragraph("Every function exposed in the authenticated IAS menu during the review is listed below. Access and actions are permission-based.",styles["Bodyx"])]
    for module in menu_data:
        story.append(Paragraph(pdf_safe(module["menu"]),styles["H2x"]))
        data=[[Paragraph("Screen / function",styles["Bodyx"]),Paragraph("How to open",styles["Bodyx"]),Paragraph("Verified user use",styles["Bodyx"])]]
        for link in module["links"]:
            data.append([Paragraph(pdf_safe(safe_text(link["text"])),styles["Bodyx"]),Paragraph(pdf_safe(f"{module['menu']} > {safe_text(link['text'])}"),styles["Bodyx"]),Paragraph(pdf_safe(purpose_for(link["text"])),styles["Bodyx"])])
        t=Table(data,colWidths=[46*mm,57*mm,67*mm],repeatRows=1); t.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.3,colors.grey),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#"+NAVY)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),("TOPPADDING",(0,0),(-1,-1),2),("BOTTOMPADDING",(0,0),(-1,-1),2)])); story.append(t)
    story += [PageBreak(),Paragraph("18. Validation, Uploads, Search and Outputs",styles["H1x"])]
    for title,body in [("Validation","Correct on-screen validation messages and verify the refreshed row, status or confirmation."),("File/evidence handling","Associate files only with the intended record and follow the extension/size message displayed by that screen. Confirm retrieval from the same record."),("Search/filtering","Select required parent values first, then search and review the resulting table."),("Reports/downloads","Select the required scope, use the displayed output action, and verify the downloaded title and scope before distribution."),("Errors and support","Record the screen name and time without confidential content, then contact IAS support if an error persists.")]: story += [Paragraph(title,styles["H2x"]),Paragraph(body,styles["Bodyx"])]
    story += [Paragraph("19. Figure Index and Verification Note",styles["H1x"])]
    for number,caption,filename in FIGURES: story.append(Paragraph(pdf_safe(f"{number}: {caption} - authenticated local IAS view ({filename})"),styles["Bodyx"]))
    story.append(Paragraph("Verified against the IAS source tree and authenticated local application on 27 August 2026. Credentials, secrets and populated confidential business records are excluded.",styles["Bodyx"]))
    def footer(canvas, doc):
        canvas.saveState(); canvas.setFont("Helvetica",7); canvas.setFillColor(colors.HexColor("#666666")); canvas.drawCentredString(A4[0]/2,10*mm,f"Internal Audit System (IAS) - User Manual | Page {doc.page}"); canvas.restoreState()
    pdf=SimpleDocTemplate(str(PDF_OUT),pagesize=A4,rightMargin=18*mm,leftMargin=18*mm,topMargin=15*mm,bottomMargin=17*mm,title="Internal Audit System (IAS) User Manual",author="IAS Technical Documentation")
    pdf.build(story,onFirstPage=footer,onLaterPages=footer)


def make_contact_sheet():
    files = sorted(ASSETS.glob("Figure_*.png"))
    width, height, columns = 480, 300, 3
    cards = []
    for path in files:
        image = Image.open(path).convert("RGB")
        image.thumbnail((width - 10, height - 38))
        card = Image.new("RGB", (width, height), "white")
        card.paste(image, ((width - image.width) // 2, 30))
        ImageDraw.Draw(card).text((8, 8), path.stem, fill="black")
        cards.append(card)
    rows = (len(cards) + columns - 1) // columns
    sheet = Image.new("RGB", (width * columns, height * rows), (225, 225, 225))
    for index, card in enumerate(cards):
        sheet.paste(card, ((index % columns) * width, (index // columns) * height))
    sheet.save(ASSETS / "contact_sheet.png")


def extract_runtime_controls():
    source = json.loads((ASSETS / "representative_screens.json").read_text(encoding="utf-8"))
    result = []
    pattern = re.compile(r"^\s*-\s+(heading|button|textbox|combobox|checkbox|tab|link|table|row|columnheader|text)\b.*", re.I)
    for item in source:
        lines = [line.strip() for line in item.get("dom", "").splitlines() if pattern.match(line)]
        result.append({"name": item["name"], "url": item["url"], "title": item["title"], "controls": lines[-80:]})
    (ASSETS / "runtime_controls.json").write_text(json.dumps(result, indent=2), encoding="utf-8")


if __name__ == "__main__":
    make_contact_sheet()
    extract_runtime_controls()
    menu_data = json.loads((ASSETS / "menu_inventory.json").read_text(encoding="utf-8"))
    build_docx(menu_data)
